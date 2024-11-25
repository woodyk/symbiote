#!/usr/bin/env python3
#
# sym_utils.py

import json
import math
import re
import os
import time
import subprocess
import magic
import hashlib
import pytesseract
import requests 
import webbrowser
import html2text
import pandas as pd
import speech_recognition as sr
import pdfplumber
from collections import Counter

from docx import Document
from datetime import datetime
from mss import mss
from urllib.parse import urlparse
from bs4 import BeautifulSoup
from io import BytesIO, StringIO
from climage import color_to_flags, color_types, convert
from PIL import Image
from ascii_magic import AsciiArt, Back
from pydub import AudioSegment
from thefuzz import fuzz
from sumy.parsers.plaintext import PlaintextParser
from sumy.nlp.tokenizers import Tokenizer
from sumy.summarizers.lsa import LsaSummarizer
from dateutil.parser import parse
from elasticsearch import Elasticsearch, exceptions
from pathlib import Path
from rich.syntax import Syntax
from rich.panel import Panel
from rich.console import Console
from rich.markdown import Markdown
console = Console()
log = console.log

import warnings
warnings.filterwarnings("ignore", category=UserWarning, module="sumy")

def clean_path(path):
    if is_url(path):
        return path
    else:
        path = os.path.expanduser(path)
        path = os.path.abspath(path)
        
        if os.path.isfile(path) or os.path.isdir(path):
            return path

def get_screenshot():
    ''' Take screenshot and return text object for all text found in the image '''
    # Screenshot storage path
    path = r'/tmp/sym_screenshot.png'

    with mss() as sct:
        monitor = {"top": 0, "left": 0, "width": 0, "height": 0}
        
        for mon in sct.monitors:
            # get furthest left point
            monitor["left"] = min(mon["left"], monitor["left"])
            # get highest point
            monitor["top"] = min(mon["top"], monitor["top"])
            # get furthest right point
            monitor["width"] = max(mon["width"]+mon["left"]-monitor["left"], monitor["width"])
            # get lowest point
            monitor["height"] = max(mon["height"]+mon["top"]-monitor["top"], monitor["height"])
        
        screenshot = sct.grab(monitor)

    img = Image.frombytes("RGB", screenshot.size, screenshot.bgra, "raw", "BGRX")
    img_gray = img.convert("L")
    img_gray.save(path)

    return path

def get_sha256(file_path):
    ''' Take in a file path and return SHA256 value for the file '''
    sha256_hash = hashlib.sha256()

    with open(file_path, "rb") as f:  # Open the file in binary mode
        for byte_block in iter(lambda: f.read(4096), b""):
            sha256_hash.update(byte_block)

    return sha256_hash.hexdigest()

def extract_exif(file_path):
    exif_data = None
    try:
        result = subprocess.run(['exiftool', '-j', file_path], stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        if result.returncode == 0:
            exif_data = json.loads(result.stdout.decode())[0]

    except Exception as e:
        log(f"Exiftool failed: {e}")

    return exif_data

def is_valid_date(date_str):
    ''' Check if an object fits the pattern of a potential date '''
    try:
        dt = parse(date_str, fuzzy=True)
        if dt.hour > 0 or dt.minute > 0:
            return True  # Time
        if dt.month > 0 and dt.day > 0 and dt.year > 0:
            return True  # Month Day Year
        return False
    except:
        return False

def remove_special_chars(values):
    if type(values) == str:
        values = re.sub(r'[^\-\.,\#A-Za-z0-9 ]+', '', values)
    elif type(values) == (str or list or tuple):
        for value in values:
            values[index(value)] = re.sub(r'[^\-\.,\#A-Za-z0-9 ]+', '', value)

    return values 

def is_source_code(text):
    score = 0

    # Look for common programming constructs
    patterns = [
        (r'\bfor\b', 1),      # for loop
        (r'\bwhile\b', 1),    # while loop
        (r'\bif\b', 1),       # if statement
        (r'\belse\b', 1),     # else statement
        (r'\bdef\b', 1),      # function definition (Python)
        (r'\bfunction\b', 1), # function definition (JavaScript)
        # Add more patterns as needed, with associated scores
    ]

    # Check for indentation
    if re.search(r'^\s', text, re.MULTILINE):
        score += 2  # Adjust score value as needed

    # Check for shebang line
    if re.search(r'^#!', text, re.MULTILINE):
        score += 3  # Adjust score value as needed

    for pattern, pattern_score in patterns:
        if re.search(pattern, text):
            score += pattern_score

    # Determine if text is likely source code based on total score
    # Adjust threshold as needed
    return score >= 5


def analyze_file(path, reindex=False):
    file_list = []
    if os.path.isdir(path):
        for root, _, files in os.walk(path):
            for file in files:
                full_path = os.path.join(root, file)
                if not os.path.isdir(full_path):
                    file_list.append(full_path)
    elif os.path.isfile(path):
        file_list.append(path)

    fcount = 0
    for file in file_list:
        fcount += 1
        doc_id = get_sha256(file)
        content = json.dumps(summarize_file(file),indent=4, sort_keys=True)
    return content

def analyze_text(text, meta=""):
    import spacy
    from spacy.lang.en.stop_words import STOP_WORDS
    from nltk.sentiment import SentimentIntensityAnalyzer
    nlp = spacy.load('en_core_web_sm')
    sia = SentimentIntensityAnalyzer()
    tokenizer = Tokenizer("english")

    try:
        text = text.decode('utf-8')
    except:
        pass

    doc = nlp(text)

    label_map = {'PERSON': 'PERSONS',
                 'NORP': 'NATIONALITIES',
                 'FAC': 'LANDMARKS',
                 'ORG': 'ORGANIZATIONS',
                 'GPE': 'LOCALITIES',
                 'LOC': 'LOCATIONS',
                 'PRODUCT': 'PRODUCTS',
                 'EVENT': 'EVENTS',
                 'WORK_OF_ART': 'ARTWORKS',
                 'LAW': 'LEGAL',
                 'LANGUAGE': 'LANGUAGES',
                 'DATE': 'DATES',
                 'TIME': 'TIMES',
                 '#PERCENT': 'PERCENTAGES',
                 'MONEY': 'CURRENCIES',
                 'QUANTITY': 'QUANTITIES',
                 '#ORDINAL': 'ORDINALS',
                 '#CARDINAL': 'CARDINALS',
                 }

    # Document container
    content = {}

    # Iterate over the entities
    ent_count = {}
    for ent in doc.ents:
        # Only include common labels
        if ent.label_ in label_map: 
            # Increment the count of the entity text for the given label
            if label_map[ent.label_] not in content:
                content[label_map[ent.label_]] = []
            elif ent.text not in content[label_map[ent.label_]] and ent.text is not (None or ""):
                clean_text = remove_special_chars(ent.text)
                content[label_map[ent.label_]].append(clean_text)

    sentiment = sia.polarity_scores(text)

    parser = PlaintextParser.from_string(text, tokenizer)
    stop_words = list(STOP_WORDS)
    summarizer = LsaSummarizer()
    summarizer.stop_words = stop_words
    summary = summarizer(parser.document, 10)
    main_idea = " ".join(str(sentence) for sentence in summary)

    content['EPOCH'] = time.time()
    #content['ADDRESSES'] = extract_address(text)
    #content.update(meta)
    content['METADATA'] = meta
    content['SENTIMENT'] = sentiment
    #content['CONTENTS'] = text
    content['SUMMARY'] = main_idea

    content['EMAILS'] = extract_email(text)
    content['WEBSITES'] = extract_url(text)
    content['PHONE_NUMBERS'] = extract_phone(text)
    content['CREDIT_CARDS'] = extract_credit_card(text)
    content['SOCIAL_SECURITY_NUMBERS'] = extract_social_securty(text)

    return content 

def summarize_text(text):
    result = analyze_text(text, meta)
    return result

def summarize_file(file_path):
    text = extract_text(file_path)

    result = analyze_text(text, meta)

    return result

def extract_dir_text(dir_path):
    file_path = cleanFilePath(file_path)

    if not os.path.isdir(dir_path):
        return None

    '''
    if os.path.isfile(".gitignore"):
        with open(os.path.join(dir_path, '.gitignore'), 'r') as f:
            gitignore = f.read()

        spec = pathspec.PathSpec.from_lines(pathspec.patterns.GitWildMatchPattern, gitignore.splitlines())
    '''

    header = str()
    content = str()
    for root, dirs, files in os.walk(dir_path):
        # Remove hidden directories from dirs so os.walk doesn't process them
        dirs[:] = [d for d in dirs if not d.startswith('.')]

        for file in files:
            # Skip hidden files
            if file.startswith('.'):
                continue

            file_path = os.path.join(root, file)
            absolute_path = clean_path(file_path)
            #if spec.match_file(file_path):
            #    continue
            file_contents = extract_text(absolute_path)
            content += f"File name: {absolute_path}\n"
            content += '\n```\n{}\n```\n'.format(file_contents)

    return content

def extract_text(file_path):
    file_path = clean_path(file_path)
    mime_type = magic.from_file(file_path, mime=True)
    try:
        content = "" 
        if mime_type.startswith('text/') or mime_type in ['application/json', 'application/xml', 'application/x-yaml', 'text/markdown']:
            with open(file_path, 'r') as f:
                content = f.read()

        elif mime_type in ['application/vnd.ms-excel', 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet']:
            content = excel_to_csv(file_path)

        elif mime_type == 'application/pdf':
            content = pdf_to_markdown(file_path)

        elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            content = word_to_markdown(file_path)

        elif mime_type.startswith('image/'):
            content = handle_image_file(file_path)

        elif mime_type.startswith('audio/'):
            content = trascrbe_audio_file(file_path)

        if len(content) > 0:
            content = content.encode('utf-8').decode('utf-8', errors='ignore')
            return content
        else:
            log(f"No content found for file: {file_path}")
            return None

    except Exception as e:
        log(f"Error reading {file_path}: {e}")
        return None

def es_connect():
    es = Elasticsearch(settings['elasticsearch'])

    if not es.ping():
        log(f'Unable to reach {settings["elasticsearch"]}')
        return None

    return es

def create_es_index(path, reindex=False):
    es = es_connect()
    
    file_list = []
    if os.path.isdir(path):
        for root, _, files in os.walk(path):
            for file in files:
                full_path = os.path.join(root, file)
                if not os.path.isdir(full_path):
                    file_list.append(full_path)
    elif os.path.isfile(path):
        file_list.append(path)

    index = settings['elasticsearch_index']

    if not es.indices.exists(index=index):
        es.indices.create(index=index)

    fcount = 0
    for file in file_list:
        fcount += 1
        if settings['debug']:
            log(f'Processing file {file}. count:{fcount}')

        doc_id = get_sha256(file)

        if not reindex:
            if es.exists(index=index, id=doc_id):
                if settings['debug']:
                    log(f"Document {doc_id} found. skipping...")
                continue

        content = summarize_file(file)

        try:
            es.index(index=index, id=doc_id, document=json.dumps(content))
        except exceptions.NotFoundError as e:
            if settings['debug']:
                log(f"Document not found: {e}")
        except exceptions.RequestError as e:
            if settings['debug']:
                log(f"Problem with the request: {e}")
        except exceptions.ConnectionError as e:
            if settings['debug']:
                log(f"Problem with the connection: {e}")
        except exceptions.ConflictError as e:
            if settings['debug']:
                log(f"Conflict occurred. Probably the document with this id already exists.")
        except exceptions.TransportError as e:
            if settings['debug']:
                log(f"General transport error: {e}")
        except Exception as e:
            if settings['debug']:
                log(f"Error: {e}")

        es.indices.refresh(index=index)

    return True

def search_index(query):
    es = es_connect()

    ret = ""

    try:
        res = es.search(index=settings['elasticsearch_index'],
                           body={
                              "track_total_hits": True,
                              "sort": [
                                {
                                  "_score": {
                                    "order": "desc"
                                  }
                                }
                              ],
                              "fields": [
                                {
                                  "field": "*",
                                  "include_unmapped": "true"
                                }
                              ],
                              "size": 10000,
                              "version": True,
                              "script_fields": {},
                              "stored_fields": [
                                "*"
                              ],
                              "runtime_mappings": {},
                              "_source": False,
                              "query": {
                                "bool": {
                                  "must": [
                                    {
                                      "query_string": {
                                        "query": query,
                                        "analyze_wildcard": True,
                                        "time_zone": "America/New_York"
                                      }
                                    }
                                  ],
                                  "filter": [],
                                  "should": [],
                                  "must_not": []
                                }
                              }
                            }
                        )

        ret = res.copy()
    except:
        log(f'Error running query: {query}')

    return ret 

def display_documents(json_data):
    # Load the JSON data
    data = json.loads(json_data)
    output = None

    display_fields = ["METADATA.SourceFile"]

    # Extract the 'hits' data
    hits = data["hits"]["hits"]

    # Create a list to store the documents
    documents = []

    # Iterate through the hits and extract the fields data
    for hit in hits:
        fields = hit["fields"]

        fields = {k: v for k, v in fields.items() if not k.endswith('.keyword')}

        documents.append(fields)

    # Create a Pandas DataFrame from the documents
    df = pd.DataFrame(documents)

    # Print the DataFrame as a table
    try:
        output = df[display_fields]
    except:
        log("No results found.")

    return output 

""""
def grep_files(es_results, search_term):
    file_list = []
    fuzzy = re.sub(r'~\d\b|\bAND\b|\bOR\b', ' ', search_term)

    regex_search = lucene_like_to_regex(search_term)
    if regex_search is False:
        return None

    for hit in es_results['hits']['hits']:
        source_file = hit["fields"]['METADATA.SourceFile'][0]
        if source_file is not None:
            file_list.append(source_file)

    text = ""
    for file_path in file_list:
        with open(file_path, 'r') as file:
            for line_no, line in enumerate(file.readlines(), start=1):
                if re.search(regex_search, line, re.I):
                    text += line
                    continue
                
                chunks = break_text(line, 1) 

                for chunk in chunks:
                    ratio = fuzz.ratio(fuzzy.lower(), chunk.lower())
                    if ratio > 50:
                        text += line
                        break
    return text
"""

def break_text(text, num_words):
    words = text.split()
    chunks = [' '.join(words[i:i + num_words]) for i in range(0, len(words), num_words)]
    return chunks

def convert_audio_to_wav(file_path):
    # Extract the file extension
    _, ext = os.path.splitext(file_path)
    ext = ext.lstrip('.')

    # Use pydub to convert to WAV
    audio = AudioSegment.from_file(file_path, format=ext)
    wav_file_path = file_path.replace(ext, 'wav')
    audio.export(wav_file_path, format='wav')

    return wav_file_path

def transcribe_audio_file(audio_file):
    recognizer = sr.Recognizer()
    _, ext = os.path.splitext(audio_file)
    text = ""
    # Convert the file to WAV if necessary
    if ext.lower() not in ['.wav', '.wave']:
        audio_file = convert_audio_to_wav(audio_file)
    try:
        with sr.AudioFile(audio_file) as source:
            audio = recognizer.record(source)

        text = recognizer.recognize_google(audio)
    except sr.UnknownValueError:
        log("Google Speech Recognition could not understand audio")
        return None
    except sr.RequestError as e:
        log(f"Could not request results from Google Speech Recognition service; {e}")
        return None

    return text

def lucene_like_to_regex(query):
    # Replace field:term to term
    single_term_regex = re.sub(r'\S+:(\S+)', r'\1', query)

    # Escape special regex characters, but leave our syntax elements
    escaped = re.sub(r'([\\.+^$[\]{}=!<>|:,\-])', r'\\\1', single_term_regex)

    # Restore escaped spaces (i.e., '\ ' to ' ')
    escaped = re.sub(r'\\ ', ' ', escaped)

    # Process grouping parentheses and quoted strings
    groups_and_quotes = re.sub(r'([()])', r'\\\1', escaped)
    groups_and_quotes = re.sub(r'"(.*?)"', r'\1', groups_and_quotes)

    # Convert wildcard queries to regex
    wildcard_regex = groups_and_quotes.replace('?', '.').replace('*', '.*')

    # Convert TO (range) queries to regex
    range_regex = re.sub(r'\[(\d+)\sTO\s(\d+)\]', lambda m: f"[{m.group(1)}-{m.group(2)}]", wildcard_regex)

    # Convert AND, OR and NOT queries to regex
    # AND operator is a bit tricky. We use positive lookaheads to emulate AND behavior in regex
    and_operator_regex = re.sub(r'(\S+)\sAND\s(\S+)', r'(?=.*\1)(?=.*\2)', range_regex)
    or_operator_regex = and_operator_regex.replace(' OR ', '|')
    not_operator_regex = or_operator_regex.replace(' NOT ', '^(?!.*')

    # Closing parentheses for each NOT operator
    final_regex = not_operator_regex.replace(' ', ').*')

    try:
        re.compile(final_regex)
        return final_regex
    except re.error:
        log(f"Invalid search term: {query}")
        return False

def scroll_content(file_path, speed=0.01):
    if os.path.isfile(file_path):
        content = extract_text(file_path)
    else:
        log(f"Path not file: {file_path}")

    for line in text.splitlines():
        print(line)
        time.sleep(speed)

def image_to_ascii(image_path):
    if is_url(image_path):
        image_path = downloadImage(image_path)

    # Get terminal size
    rows, columns = os.popen('stty size', 'r').read().split()
    term_width = int(columns)

    # Calculate 70% of the terminal width
    image_width = int(term_width * 0.7)

    # Display the ASCII art, scaled to fit 70% of the terminal width
    image = convert(image_path, width=image_width, is_unicode=True, **color_to_flags(color_types.color256))
    remove_file(file_path)

    # Calculate padding to center the image
    padding = (term_width - image_width) // 2
    padded_image = "\n".join([" " * padding + line for line in image.split("\n")])

    return padded_image

def remove_file(file_path):
    file_path = clean_path(file_path)
    try:
        os.remove(file_path)
        log(f"File removed: {file_path}")
        return True
    except FileNotFoundError:
        log(f"File not found: {file_path}")
        return False
    except PermissionError:
        log(f"Permission denied: {file_path}")
        return False
    except Exception as e:
        log(f"Error removing file: {e}")
        return False

def downloadImage(url):
    if is_image(url):
        filename = os.path.basename(urlparse(url).path)
        save_path = os.path.join('/tmp/', filename)

        response = requests.get(url, stream=True)
        response.raise_for_status()

        with open(save_path, 'wb') as file:
            for chunk in response.iter_content(chunk_size=8192):
                file.write(chunk)

        return clean_path(save_path)
    else:
        log(f"Unable to pull image from {url}")
        return None

def is_image(file_path_or_url):
    try:
        if is_url(file_path_or_url):
            response = requests.head(file_path_or_url, allow_redirects=True)
            content_type = response.headers.get("Content-Type", "").lower()
            if content_type.startswith("image/"):
                return True
        else:
            mime = magic.from_file(file_path_or_url, mime=True)
            if mime.startswith("image/"):
                return true
    except Exception as e:
        return False

def is_url(path):
    url_pattern = re.compile(
        r'^(?:http|ftp)s?://',
        re.IGNORECASE)
    if re.match(url_pattern, path):
        return True

    return None

def pdf_to_markdown(pdf_path):
    """
    Converts a PDF to Markdown, including text, tables, and images.
    """
    markdown_content = ""

    def table_to_markdown(table):
        """Converts a PDF table to Markdown format."""
        table_md = ""
        for row in table:
            table_md += "| " + " | ".join(row) + " |\n"
        if table and len(table[0]) > 0:
            table_md += "|---" * len(table[0]) + "|\n"
        return table_md

    def save_image(image_data, page_num, image_num):
        """Saves an image and returns the Markdown link to it."""
        image_dat = ""
        image_path = f"/tmp/page-{page_num}-image-{image_num}.png"
        with open(image_path, "wb") as img_file:
            img_file.write(image_data)

        image_md = handle_image_file(image_path)
        image_dat += f"![Image {image_num}](./{image_path})\n\n"
        if image_md:
            image_dat += image_md

    with pdfplumber.open(pdf_path) as pdf:
        # Extract metadata
        metadata = pdf.metadata
        if metadata:
            markdown_content += "## Metadata\n"
            for key, value in metadata.items():
                markdown_content += f"- **{key}**: {value}\n"
            markdown_content += "\n"

        # Process each page
        if pdf.pages:
            for page_num, page in enumerate(pdf.pages, start=1):
                markdown_content += f"\n\n## Page {page_num}\n\n"

                # Extract text
                try:
                    text = page.extract_text()
                    if text:
                        markdown_content += text + "\n\n"
                    else:
                        markdown_content += "No text found on this page.\n\n"
                except Exception as e:
                    log(f"Error processing pdf: {pdf_path}\n{e}")

                # Extract tables
                try:
                    tables = page.extract_tables()
                    if tables:
                        for table_num, table in enumerate(tables, start=1):
                            markdown_content += f"### Table {table_num}\n\n"
                            markdown_content += table_to_markdown(table)
                    else:
                        markdown_content += "No tables found on this page.\n\n"
                except Exception as e:
                    log(f"Error processing pdf: {pdf_path}\n{e}")

                # Extract images
                try:
                    if page.images:
                        for image_num, image in enumerate(page.images, start=1):
                            if "data" in image:
                                image_data = image["data"]
                                image_md = save_image(image_data, page_num, image_num)
                                markdown_content += f"{image_md}\n\n"
                    else:
                        markdown_content += "No images found on this page.\n\n"
                except Exception as e:
                    log(f"Error processing pdf: {pdf_path}\n{e}")

        return markdown_content

def word_to_markdown(file_path):
    """
    Converts a Word document (.docx) to Markdown format, including headings, text, tables, and lists.
    """
    file_path = clean_path(file_path)
    # Load the Word document
    doc = Document(file_path)
    markdown_content = ""

    # Helper function to process tables
    def table_to_markdown(table):
        table_md = ""
        # Iterate through table rows
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            table_md += "| " + " | ".join(cells) + " |\n"
        if len(table.rows) > 0:
            # Add separator for table header
            table_md += "|---" * len(table.rows[0].cells) + "|\n"
        return table_md

    # Helper function to process lists (ordered and unordered)
    def list_to_markdown(paragraph, indent_level):
        text = paragraph.text.strip()
        if paragraph.style.name.startswith("List Bullet"):
            return f"{'  ' * indent_level}- {text}\n"
        elif paragraph.style.name.startswith("List Number"):
            return f"{'  ' * indent_level}1. {text}\n"
        return ""

    # Iterate through all elements in the document
    try:
        for paragraph in doc.paragraphs:
            # Handle headings
            if paragraph.style.name.startswith("Heading"):
                heading_level = int(paragraph.style.name[-1])  # Get heading level (1-6)
                markdown_content += f"{'#' * heading_level} {paragraph.text.strip()}\n\n"
            # Handle lists
            elif paragraph.style.name.startswith("List"):
                markdown_content += list_to_markdown(paragraph, indent_level=1)
            # Handle normal text
            elif paragraph.text.strip():
                markdown_content += f"{paragraph.text.strip()}\n\n"
    except Exception as e:
        log(f"Error reading {file_path}:\n{e}")
        return None

    # Add tables
    try:
        for table in doc.tables:
            markdown_content += "\n" + table_to_markdown(table) + "\n"
    except Exception as e:             
        log(f"Error readingt {file_path}:\n{e}")
        return None  

    return markdown_content

def excel_to_csv(file_path):
    """
    Converts an Excel file to CSV format.
    """
    file_path = clean_path(file_path)
    csv_content = ""
    try:
        df = pd.read_excel(file_path)
        ouput = StringIO()
        df.to_csv(file_path, index=False)
        csv_content = output.getvalue()
    except Exception as e:
        log(f"Failed to convert Excel to CSV: {e}")

    return csv_content

def handle_image_file(file_path):
    file_path = clean_path(file_path)
    mime_type = magic.from_file(file_path, mime=True)
    try:
        with Image.open(file_path) as img:
            # Extract metadata
            metadata = {
                "format": img.format,
                "size": img.size,  # (width, height)
                "mode": img.mode,  # Color mode
                "mime_type": mime_type
            }

            # Perform OCR to extract text
            extracted_text = pytesseract.image_to_string(img)

            # Create a well-formatted Markdown document
            markdown_content = StringIO()
            markdown_content.write("# Image Analysis Report\n\n")
            markdown_content.write(f"## Metadata\n")
            for key, value in metadata.items():
                markdown_content.write(f"- **{key.capitalize()}**: {value}\n")

            markdown_content.write("\n## Extracted Text\n")
            if extracted_text.strip():
                markdown_content.write("```\n")
                markdown_content.write(extracted_text)
                markdown_content.write("\n```\n")
            else:
                markdown_content.write("No text could be extracted from the image.\n")

            log(f"Processed image file and generated Markdown content.")
            return markdown_content.getvalue()
    except Exception as e:
        log(f"Failed to process image: {e}")

    return None

def extract_metadata(file_path, output=None):
    """
    Extracts information from a file of unknown or unsupported type.
    """
    file_path = clean_path(file_path)
    file_stats = os.stat(file_path)
    creation_time = datetime.fromtimestamp(file_stats.st_ctime)
    modified_time = datetime.fromtimestamp(file_stats.st_mtime)
    file_info = {
        "file_path": str(file_path),
        "file_size": str(file_stats.st_size),
        "creation_time": str(creation_time),
        "modification_time": str(modified_time), 
        "permissions": oct(file_stats.st_mode & 0o777),
        "mime_type": str(magic.from_file(file_path, mime=True)),
        "hashes": {},
        "readable_strings": [],
        "magic_numbers": None,
        "embedded_urls": [],
        "entropy": None,
        "exif": {},
    }

    exif_data = extract_exif(file_path) 
    for key, value in exif_data.items():
        if key not in file_info:
            file_info["exif"][key] = value 

    def calculate_entropy(data):
        """Calculate Shannon entropy to assess randomness in the file."""
        if not data:
            return 0
        counter = Counter(data)
        length = len(data)
        entropy = -sum((count / length) * math.log2(count / length) for count in counter.values())
        return str(entropy)

    def extract_strings(data):
        """Extract readable ASCII and Unicode strings."""
        ascii_regex = re.compile(rb'[ -~]{4,}')  # ASCII strings of length >= 4
        unicode_regex = re.compile(rb'(?:[\x20-\x7E][\x00]){4,}')  # Unicode UTF-16 strings
        strings = []
        strings.extend(match.decode('ascii') for match in ascii_regex.findall(data))
        strings.extend(match.decode('utf-16') for match in unicode_regex.findall(data))
        return strings

    def find_binary_urls(data):
        """Find embedded URLs in the file."""
        url_regex = re.compile(rb'https?://[^\s]+')  # Regex to match URLs
        urls = []
        for match in url_regex.findall(data):
            try:
                urls.append(match.decode('ascii'))  # Try decoding with ASCII
            except UnicodeDecodeError:
                urls.append(match.decode('utf-8', errors='ignore'))  # Fall back to UTF-8 with error handling
            urls = list(set(urls))
        return urls

    # Read the file as binary
    with open(file_path, 'rb') as file:
        binary_data = file.read()
        # Compute hashes
        file_info["hashes"]["sha256"] = hashlib.sha256(binary_data).hexdigest()
        file_info["hashes"]["md5"] = hashlib.md5(binary_data).hexdigest()
        # Extract readable strings
        file_info["readable_strings"] = extract_strings(binary_data)
        # Extract embedded URLs
        file_info["embedded_urls"] = find_binary_urls(binary_data)
        # Calculate entropy
        file_info["entropy"] = calculate_entropy(binary_data)
        # Extract magic numbers (first 4 bytes)
        file_info["magic_numbers"] = binary_data[:4].hex()

    if output == "json":
        return json.dumps(file_info, indent=4)
    elif output == "markdown":
        # Generate a Markdown report
        markdown_content = f"# Catch-All Analysis Report\n\n"
        markdown_content += f"**File Path**: {file_info['file_path']}\n"
        markdown_content += f"**File Size**: {file_info['file_size']} bytes\n"
        markdown_content += f"**MIME Type**: {file_info['mime_type']}\n"
        markdown_content += f"**Permissions**: {file_info['permissions']}\n"
        markdown_content += f"**Creation Time**: {file_info['creation_time']}\n"
        markdown_content += f"**Modification Time**: {file_info['modification_time']}\n"
        markdown_content += f"**SHA-256 Hash**: {file_info['hashes']['sha256']}\n"
        markdown_content += f"**MD5 Hash**: {file_info['hashes']['md5']}\n"
        markdown_content += f"**Magic Numbers**: {file_info['magic_numbers']}\n"
        markdown_content += f"**Entropy**: {file_info['entropy']:.4f}\n\n"

        if "exif" in file_info:
            markdown_content += "## Exif Data (Exiftool)\n"
            for key, value in file_info["exif"]:
                markdown_content += f"**{key}***: {value}\n"

        markdown_content += "## Readable Strings\n"
        markdown_content += "\n".join(f"- {string}" for string in file_info['readable_strings'][:10])  # Show first 10
        markdown_content += "\n\n"

        markdown_content += "## Embedded URLs\n"
        markdown_content += "\n".join(f"- {url}" for url in file_info['embedded_urls'])
        markdown_content += "\n\n"

        return markdown_content
    else:
        return file_info

def extract_pii(text):
    extract_url(text)
    extract_email(text)
    extract_phone(text)
    extract_address(text)
    extract_medical(text)
    extract_credit_card(text)
    extract_social_securty(text)

def extract_email(text):
    ''' Check text for the pattern of an e-mail address '''
    email_pattern = r"[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+"
    matches = re.findall(email_pattern, text)
    
    clean = clean_matches(matches)

    return clean

def clean_matches(matches):
    clean = []
    for match in matches:
        if type(match) == tuple or type(match) == list:
            for entry in match:
                if entry is not None and not "" and entry not in clean:
                    clean.append(entry)
        elif type(match) == str:
            clean.append(match)
        else:
            pass

    return clean

def extract_credit_card(text):
    card_pattern = r'\b(?:\d[ -]*?){13,16}\b'
    matches = re.findall(card_pattern, text)

    clean = clean_matches(matches)
    clean = remove_special_chars(clean)

    return clean

def extract_social_securty(text):
    ss_pattern = r'\b\d{3}-?\d{2}-?\d{4}\b'
    matches = re.findall(ss_pattern, text)

    clean = clean_matches(matches)
    clean = remove_special_chars(clean)

    return clean

def extract_medical(text):
    import spacy
    nlpm = spacy.load("en_core_sci_sm")

    doc = nlpm(text)

    for ent in doc.ents:
        log(ent.label_, ent.text)

def extract_url(text):
    #url_pattern = r"(?:http[s]?:\/\/)?(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+"
    #url_pattern = r'(?:http[s]?:\/\/)?[\w.-]+(?:\.[\w\.-]+)+[\w\-\._~:/?#[\]@!\$&\'\(\)\*\+,;=.]+'
    url_pattern = r"(?i)\b((?:https?://|www\d{0,3}[.]|[a-z0-9.\-]+[.][a-z]{2,4}/)(?:[^\s()<>]+|\(([^\s()<>]+|(\([^\s()<>]+\)))*\))+(?:\(([^\s()<>]+|(\([^\s()<>]+\)))*\)|[^\s`!()\[\]{};:'\".,<>?«»“”‘’]))"
    matches = re.findall(url_pattern, text)

    clean = clean_matches(matches)

    return clean

def extract_phone(text):
    phone_number_pattern = r"\b\d{10,11}\b|\b(\(\d{3}\)\s*\d{3}[-\.\s]??\d{4}|\d{3}[-\.\s]??\d{3}[-\.\s]??\d{4}|\d{3}[-\.\s]??\d{4})\b"
    matches = re.findall(phone_number_pattern, text)

    clean = clean_matches(matches)
    clean = remove_special_chars(clean)

    return clean

def extract_address(text):
    #from postal.parser import parse_address
    components = [ 'house_number', 'road', 'postcode', 'city', 'state' ]
    parsed_address = parse_address(text)
    addresses = []
    address = {}

    for component in parsed_address:
        if component[1] in address:
            addresses.append(address)
            address = {}
        if component[1] in components:
            address[component[1]] = remove_special_chars(component[0])

    return addresses

def is_file_path(string):
    """
    Check if a given string matches the pattern of a file path in an OS-agnostic manner.
    
    :param string: The string to check.
    :return: True if the string matches a file path pattern, False otherwise.
    """
    if not isinstance(string, str):
        return False
    
    # OS-agnostic path separators
    path_sep_pattern = re.escape(os.sep)  # For current OS
    alt_sep_pattern = re.escape(os.altsep) if os.altsep else ""  # For alternative sep if exists
    path_separators = f"[{path_sep_pattern}{alt_sep_pattern}]"
    
    # File path regex pattern
    # Matches strings like "folder/file.txt", "C:/folder/file", "/etc/config"
    path_pattern = rf"(^\.?{path_separators}|^[a-zA-Z]:{path_separators}|^{path_separators})[^:*?\"<>|\r\n]+"
    
    return re.match(path_pattern, string) is not None
